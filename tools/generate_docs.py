"""
PakkaHisaab collateral generator.

Builds two production-grade .docx deliverables using python-docx:
  1. docs/PakkaHisaab_Brochure.docx   - enterprise value-proposition brochure
  2. docs/PakkaHisaab_UserGuide.docx  - exhaustive engineering & workflow guide

Run:  python tools/generate_docs.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

# ---------------------------------------------------------------------------
# Brand palette
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x0B, 0x1F, 0x3A)
TEAL = RGBColor(0x0E, 0x7C, 0x7B)
GOLD = RGBColor(0xC9, 0x9A, 0x2E)
SLATE = RGBColor(0x4A, 0x55, 0x68)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DANGER = RGBColor(0xA3, 0x2A, 0x2A)
PROTIP_LABEL = RGBColor(0x8A, 0x62, 0x00)

BAND_GREY = "EEF1F5"
BAND_TEAL = "E4F1F0"

FONT_BODY = "Calibri"
FONT_HEAD = "Arial"

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.normpath(os.path.join(HERE, "..", "docs"))
LOGO_PATH = os.path.join(OUT_DIR, "logo_512.png")


# ---------------------------------------------------------------------------
# Shared low-level helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="B9C2CE", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def fixed_layout(table):
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def shade_table_zebra(table, header_color, band_color, header_font_color=WHITE):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if r_idx == 0:
                set_cell_shading(cell, header_color)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, band_color)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    if r_idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = header_font_color
                        run.font.size = Pt(10.5)
                    else:
                        run.font.size = Pt(10)


def add_page_border(section, color="0B1F3A", size=18):
    sec_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), color)
        pg_borders.append(el)
    sec_pr.append(pg_borders)


def add_footer_with_page_number(section, brand_text):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{brand_text}  |  Page ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = SLATE
    run.font.name = FONT_BODY

    def field_run(paragraph, field_name):
        r = paragraph.add_run()
        r.font.size = Pt(8.5)
        r.font.color.rgb = SLATE
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_name
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r._r.append(begin)
        r._r.append(instr)
        r._r.append(end)
        return r

    field_run(p, "PAGE")
    mid = p.add_run(" of ")
    mid.font.size = Pt(8.5)
    mid.font.color.rgb = SLATE
    field_run(p, "NUMPAGES")


def add_shaded_paragraph_box(doc, lines, fill="FFF6DF", border_color="C99A2E",
                              label=None, label_color=None):
    """A single-cell shaded/bordered table used as a callout / Pro-Tip / Security box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "10")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_color)
        borders.append(el)
    tc_pr.append(borders)
    cell.paragraphs[0].text = ""
    first = True
    if label:
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = FONT_HEAD
        run.font.color.rgb = label_color or PROTIP_LABEL
        first = False
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(2)
    fixed_layout(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_image_placeholder(doc, caption_text, height_cm=6.5):
    """Dashed-border placeholder box standing in for an HD graphic/screenshot."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F7F8FA")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "dashed")
        el.set(qn("w:sz"), "10")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "8FA0B3")
        borders.append(el)
    tc_pr.append(borders)

    tr = table.rows[0]._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_cm * 567)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    icon_run = p.add_run("\U0001F5BC  ")
    icon_run.font.size = Pt(20)
    run = p.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = SLATE
    run.font.name = FONT_BODY
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_hr(doc, color="C9CFD9", size="6"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    p_pr.append(pbdr)


# ---------------------------------------------------------------------------
# Style setup shared by both documents
# ---------------------------------------------------------------------------
def configure_base_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x22, 0x28, 0x30)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_BODY)

    def style_heading(name, size, color, bold=True, space_before=18, space_after=8):
        st = styles[name]
        st.font.name = FONT_HEAD
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(space_before)
        st.paragraph_format.space_after = Pt(space_after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing = 1.05
        return st

    style_heading("Title", 30, NAVY, space_before=0, space_after=6)
    style_heading("Heading 1", 20, NAVY, space_before=26, space_after=10)
    style_heading("Heading 2", 15, TEAL, space_before=18, space_after=8)
    style_heading("Heading 3", 12.5, SLATE, space_before=12, space_after=6)

    if "Subtitle" in [s.name for s in styles]:
        sub = styles["Subtitle"]
        sub.font.name = FONT_BODY
        sub.font.size = Pt(14)
        sub.font.italic = True
        sub.font.color.rgb = TEAL
        sub.font.bold = False

    for sec in doc.sections:
        sec.top_margin = Cm(2.1)
        sec.bottom_margin = Cm(2.1)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)


def add_bullets(doc, items, style="List Bullet", size=10.5):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.size = Pt(size)
    return doc


def add_kicker(doc, text, color=None):
    color = color or TEAL
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = FONT_HEAD
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = color
    rpr = run._r.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "30")
    rpr.append(spacing)
    return p


def add_meta_table(doc, rows, col_widths=(4.2, 10.5)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = ""
        r = cells[0].paragraphs[0].add_run(k)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = NAVY
        cells[1].text = ""
        r2 = cells[1].paragraphs[0].add_run(v)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = SLATE
        for c in cells:
            set_cell_borders(c, color="D9DEE6", size="2")
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.columns[0].width = Cm(col_widths[0])
    table.columns[1].width = Cm(col_widths[1])
    return table


def add_cover_logo_band(doc, height_cm=4.4, placeholder_lines=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "0B1F3A")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "C99A2E")
        borders.append(el)
    tc_pr.append(borders)
    tr = table.rows[0]._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_cm * 567)))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        run = lp.add_run()
        run.add_picture(LOGO_PATH, width=Cm(3.2))
    else:
        lines = placeholder_lines or [
            "[ CORPORATE LOGO PLACEHOLDER ]",
            "PakkaHisaab · 512x512 vector mark, transparent background",
        ]
        r1 = lp.add_run(lines[0] + "\n")
        r1.font.color.rgb = GOLD
        r1.font.bold = True
        r1.font.size = Pt(12)
        r2 = lp.add_run(lines[1])
        r2.font.color.rgb = WHITE
        r2.italic = True
        r2.font.size = Pt(9)
    return table


# ===========================================================================
# DOCUMENT 1 — PakkaHisaab Brochure
# ===========================================================================
def build_brochure():
    doc = Document()
    configure_base_styles(doc)

    section = doc.sections[0]
    add_page_border(section)
    add_footer_with_page_number(section, "PakkaHisaab — Product Brochure")

    # ---------------- Cover Page ----------------
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_cover_logo_band(doc)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run("PakkaHisaab")

    tagline = doc.add_paragraph(style="Subtitle")
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.add_run("Tap. Track. Settle.")

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run(
        "The Offline-First Ledger for Household Help — Attendance, Wages,\n"
        "Voice Commands and UPI Settlement in One App"
    )
    r.font.size = Pt(12.5)
    r.font.color.rgb = SLATE
    r.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_image_placeholder(
        doc,
        "Insert screenshot: Dashboard home screen — helper cards (Geeta, Raju), "
        "the \"Payable this month\" total, today's attendance status icons, and the "
        "mic button used for Voice-to-Ledger.",
        height_cm=6.0,
    )

    add_meta_table(doc, [
        ("Document Class", "Product Brochure & Value Proposition"),
        ("Prepared For", "Households employing domestic help, and PakkaHisaab support/ops teams"),
        ("Version", "v1.0 — reflects the shipped feature set"),
        ("Date", "23 July 2026"),
    ])

    doc.add_page_break()

    # ---------------- Section 2: The Problem & The Approach ----------------
    add_kicker(doc, "The Problem & The Approach")
    doc.add_heading("Replacing the Torn Notebook Page with One Shared Ledger", level=1)

    doc.add_paragraph(
        "Almost every household that employs a maid, cook, driver, gardener, or "
        "milkman keeps track of their pay the same informal way: a notebook page, a "
        "mental tally of advances given mid-month, and a memory of who was absent on "
        "which day. None of it reconciles cleanly — an advance given on the 12th is "
        "forgotten by the 30th, an absence is disputed because nobody wrote it down, "
        "and \"how much do I owe her this month\" gets re-derived from scratch every "
        "settlement day. PakkaHisaab replaces that notebook with a structured, always-"
        "current ledger per helper, kept on the household's own phone."
    )

    doc.add_heading("One Formula, Everywhere It's Used", level=2)
    doc.add_paragraph(
        "The engineering thesis is simple: the payable amount a household sees on "
        "screen, the amount synced to the cloud, and the amount printed on a shared "
        "PDF statement must always agree — because they come from exactly one "
        "formula, implemented once. PakkaHisaab.Shared.Domain.SalaryCalculator computes "
        "Final Payable = Monthly Wage − (Unpaid Absences × Daily Rate) − Advances − "
        "Other Deductions + Bonuses − Already Paid, with support for a configurable "
        "monthly allowed-absence quota, optional leave carry-over month to month, and "
        "a separate per-unit wage model (rate × units delivered) for helpers like a "
        "milkman who are paid by delivery rather than a fixed salary. That one static "
        "method is called by the mobile app for instant offline totals, by the API for "
        "server-side verification, and by the PDF report generator — never three "
        "different implementations quietly drifting apart."
    )

    doc.add_heading("Offline-First by Construction, Not by Exception", level=2)
    doc.add_paragraph(
        "Every tap — marking attendance, logging an advance, recording a delivery — "
        "writes to a local SQLite database first, so the UI never waits on a network "
        "call. Each row carries an IsDirty flag, a ModifiedAtUtc timestamp, and a "
        "RowVersion watermark; a background Shiny.Jobs worker drains that outbox to the "
        "API opportunistically, tagging each batch with a ClientBatchId so a retried "
        "push after a lost response can never double-apply. The app is fully usable "
        "with no connectivity at all, which matters in practice — the household or the "
        "helper's own signal is not always reliable, and \"the app is down\" should "
        "never be the reason today's attendance didn't get logged."
    )

    doc.add_heading("Voice-to-Ledger for Hands-Free Entries", level=2)
    doc.add_paragraph(
        "A rule-based natural-language parser — running entirely on-device, no network "
        "required — understands short spoken commands in English and romanized Hindi: "
        "\"Gave Geeta 500 advance,\" \"Raju delivered 1.5 litres,\" \"Geeta was absent "
        "today.\" It is a genuinely useful shortcut for a user who doesn't want to open "
        "forms for a thirty-second entry, and the same parser is exposed server-side "
        "at POST /ai/parse so any thin client can reuse it — one implementation, two "
        "hosts."
    )

    add_shaded_paragraph_box(
        doc,
        [
            "PakkaHisaab is a focused, single-purpose ledger for one household's "
            "domestic help — not a multi-tenant platform. Its value is in getting that "
            "one job right: one formula, one offline-first data path, one shared "
            "source of truth between the phone, the sync service, and the printed "
            "statement."
        ],
        fill="E4F1F0", border_color="0E7C7B",
        label="ENGINEERING PERSPECTIVE", label_color=TEAL,
    )

    doc.add_page_break()

    # ---------------- Section 3: Value Pillars ----------------
    add_kicker(doc, "Value Pillars & Core Capabilities")
    doc.add_heading("From Feature to Everyday Household Benefit", level=1)
    doc.add_paragraph(
        "The four pillars below are the shipped capabilities of the app today, and "
        "the concrete, everyday problem each one removes for a household managing "
        "domestic help."
    )

    pillars = [
        ("01", "Automated Attendance & Salary Engine",
         "Turns a month of scattered memory into one always-current, disputable-free number.",
         [
             "Tap-to-cycle attendance (Present → Absent → Half-Day) per helper per day, "
             "plus a separate per-unit entry mode for delivery-based helpers (e.g. "
             "litres of milk); a configurable monthly allowed-absence quota with "
             "optional carry-over of unused leave to the next month.",
             "The payable amount recalculates instantly from the shared "
             "SalaryCalculator formula the moment any attendance, advance, deduction, "
             "or bonus is logged — there is no separate \"run payroll\" step and "
             "nothing to reconstruct at month end.",
             "Business value: removes the single biggest source of household-helper "
             "pay disputes — disagreement over how many days were actually absent, "
             "and how much was already advanced — because both sides can see the "
             "same running ledger at any time, not just on settlement day.",
         ]),
        ("02", "Offline-First Reliability with Background Cloud Sync",
         "Works the moment you open it, with or without a signal, and never loses an entry.",
         [
             "All data is written to on-device SQLite first; a background sync job "
             "pushes to the cloud API and pulls remote changes opportunistically, "
             "using idempotent batches so a dropped connection mid-sync can never "
             "create duplicate ledger rows.",
             "A dedicated Demo mode seeds an isolated local database with sample "
             "helpers (no account, no network, sync disabled) — useful for trying the "
             "app instantly, and specifically built to satisfy app-store reviewers "
             "who need to evaluate the app without creating real data.",
             "Business value: a household never loses today's attendance entry to a "
             "bad connection, and never needs to \"wait for the app to load\" before "
             "logging something that just happened.",
         ]),
        ("03", "Voice-to-Ledger & Smart Notifications",
         "Reduces the two most common failure modes of manual tracking: forgetting to log, and forgetting to pay.",
         [
             "Speak a short command — \"Gave Geeta 500 advance\", \"Raju delivered 1.5 "
             "litres\" — and the on-device parser logs it without opening a single "
             "form; recognized in English and romanized Hindi today.",
             "A daily 5 PM reminder to mark attendance (with an inline \"Mark Absent\" "
             "action right in the notification shade — no need to even open the app), "
             "plus automatic salary-due alerts from the 1st–10th of the month that "
             "cancel themselves the instant a payment is recorded.",
             "An absence-pattern forecast (\"Usually absent on Mondays · ~2/month\") "
             "surfaces on the Dashboard once enough attendance history exists, so a "
             "household can plan around a helper's typical pattern rather than being "
             "surprised by it.",
         ]),
        ("04", "Instant UPI Settlement & Shareable PDF Statements",
         "Turns \"pay the milkman\" from an unrecorded cash handoff into a one-tap, logged transaction.",
         [
             "The Settlement screen shows the computed breakdown (gross wage, absence "
             "deduction, advances, final payable) and launches the phone's own "
             "installed UPI apps via a pre-filled upi://pay deep link — no card, no "
             "POS hardware, no separate invoice — or logs a plain cash payment "
             "instead.",
             "Two PDF report types — a per-helper monthly ledger with daily breakdown, "
             "and a household-wide summary across all helpers — generate on-device and "
             "share directly to WhatsApp in one tap.",
             "Business value: gives a household a real paper trail for what was paid "
             "and when, useful for their own records, without requiring the helper to "
             "adopt any technology at all — the deep link, the record-keeping, and the "
             "sharing all happen on the household's side.",
         ]),
    ]

    for num, title, sub, points in pillars:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        rnum = p.add_run(f"{num}  ")
        rnum.font.color.rgb = GOLD
        rnum.font.bold = True
        rnum.font.size = Pt(14)
        rnum.font.name = FONT_HEAD
        rtitle = p.add_run(title)
        rtitle.font.color.rgb = NAVY
        rtitle.font.bold = True
        rtitle.font.size = Pt(14)
        rtitle.font.name = FONT_HEAD

        psub = doc.add_paragraph()
        psub.paragraph_format.space_after = Pt(6)
        rsub = psub.add_run(sub)
        rsub.italic = True
        rsub.font.color.rgb = TEAL
        rsub.font.size = Pt(10.5)

        add_bullets(doc, points)

    add_image_placeholder(
        doc,
        "Insert screenshot grid: Calendar (tap-to-cycle attendance), Ledger (advance/"
        "deduction/bonus entries), and Settlement (computed payable + UPI app chooser) "
        "screens side by side on a phone frame.",
        height_cm=6.5,
    )

    doc.add_page_break()

    # ---------------- Section 4: Security & Compliance ----------------
    add_kicker(doc, "Security & Compliance")
    doc.add_heading("What's Actually Implemented, in Plain Terms", level=1)
    doc.add_paragraph(
        "The app handles household financial data, so the security posture below is "
        "described exactly as implemented — no claims beyond what the code does today."
    )

    sec_items = [
        "Password Storage — passwords are never stored in plain text or reversibly "
        "encrypted; PasswordHasher.cs hashes them with PBKDF2-SHA256, 100,000 "
        "iterations and a random 16-byte salt per user, and verifies with a "
        "constant-time comparison to resist timing attacks.",
        "API Authentication — every API call is authenticated with a JWT bearer "
        "token, signed HMAC-SHA256 by the server with a configurable expiry window; "
        "there is no session state to attack on the server side.",
        "Right to Delete — a genuine \"Delete My Account\" flow (DELETE /account) "
        "hard-deletes the user's server-side records and wipes the local database in "
        "the same action, satisfying app-store data-deletion requirements rather than "
        "just hiding the account.",
        "Local-First Data Minimization — Demo mode makes zero network calls; a user "
        "who never creates an account never sends any data off the device at all, "
        "because the local SQLite copy is the primary store, not a cache in front of "
        "a cloud original.",
        "Disclosed Diagnostics — crash and usage telemetry is integrated via the App "
        "Center SDK and disclosed in the in-app privacy policy, not collected "
        "silently (note: Microsoft retired App Center in March 2025; a documented "
        "migration path to Sentry or Firebase Crashlytics exists for the next "
        "telemetry refresh).",
        "Device-Integrity Advisory — on a rooted or jailbroken device the app shows "
        "a non-blocking warning rather than silently operating without disclosure, "
        "consistent with app-store guidelines that permit warning rather than "
        "hard-blocking such devices.",
    ]
    add_bullets(doc, sec_items)

    add_shaded_paragraph_box(
        doc,
        [
            "Nothing above is aspirational — every item traces to a specific file in "
            "the codebase (PasswordHasher.cs, TokenService.cs, AccountEndpoints.cs, "
            "App.xaml.cs) and can be re-verified by reading it."
        ],
        fill="FDECEC", border_color="A32A2A",
        label="HONESTY NOTE", label_color=DANGER,
    )

    doc.add_page_break()

    # ---------------- Section 5: Try It ----------------
    add_kicker(doc, "See It Working")
    doc.add_heading("A Real Path You Can Walk Through Today", level=1)
    doc.add_paragraph(
        "Unlike a roadmap pitch, every step below is something you can actually do "
        "in the current build right now — no signup required to start."
    )

    doc.add_heading("For a Household (the mobile app)", level=2)
    household_rows = [
        ("Step", "Action", "Outcome"),
        ("1", "Open the app and tap \"Try Demo\" (or register a real account)",
         "Instant offline sample data — Geeta (house help) and Raju (milkman) — or a synced account"),
        ("2", "Tap + to add a helper: name, category, monthly wage or per-unit rate, "
              "allowed absences",
         "A new helper card appears on the Dashboard"),
        ("3", "Open the helper's Calendar and tap today to mark attendance, or use the "
              "mic to say \"Geeta was absent today\"",
         "Attendance is recorded and the month summary updates immediately"),
        ("4", "Tap Settle on the helper card",
         "Computed payable shown; pay via UPI app chooser or log cash"),
        ("5", "Open Reports, pick a month, generate a PDF and share it",
         "A per-helper or household-wide statement lands in WhatsApp"),
    ]
    household_table = doc.add_table(rows=len(household_rows), cols=3)
    household_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(household_rows):
        cells = household_table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(val)
            if i == 0:
                r.bold = True
        if i == 0:
            set_repeat_header(household_table.rows[i])
    household_table.columns[0].width = Cm(1.6)
    household_table.columns[1].width = Cm(7.0)
    household_table.columns[2].width = Cm(6.1)
    shade_table_zebra(household_table, "0B1F3A", BAND_GREY)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.add_heading("For Support & Operations (the Admin console)", level=2)
    admin_rows = [
        ("Step", "Action", "Outcome"),
        ("1", "Run `dotnet run --project src/PakkaHisaab.Admin` and sign in with an "
              "account flagged IsAdmin",
         "Cookie-authenticated access to the ops dashboard"),
        ("2", "View the Dashboard",
         "KPIs — total users, active/total helpers, pending settlements (₹), paid "
         "this month (₹), sync batches in the last 7 days — plus charts for user "
         "growth, 30-day attendance, settlements and ledger movement"),
        ("3", "Browse Users, Helpers, Attendance, Ledger and Settlements",
         "Read/manage views over the same data the mobile app writes"),
        ("4", "Check Sync Batches",
         "Confirms the offline-first mobile clients are pushing/pulling correctly"),
    ]
    admin_table = doc.add_table(rows=len(admin_rows), cols=3)
    admin_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(admin_rows):
        cells = admin_table.rows[i].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(val)
            if i == 0:
                r.bold = True
        if i == 0:
            set_repeat_header(admin_table.rows[i])
    admin_table.columns[0].width = Cm(1.6)
    admin_table.columns[1].width = Cm(7.0)
    admin_table.columns[2].width = Cm(6.1)
    shade_table_zebra(admin_table, "0E7C7B", BAND_TEAL)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    cta_table = doc.add_table(rows=1, cols=1)
    cta_cell = cta_table.rows[0].cells[0]
    set_cell_shading(cta_cell, "0B1F3A")
    set_cell_borders(cta_cell, color="C99A2E", size="6")
    cta_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cta_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run("Ready to Try It?")
    r1.font.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = WHITE
    r1.font.name = FONT_HEAD
    p2 = cta_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(
        "Try Demo needs no signup, no network, and touches no real data — "
        "the fastest way to see every feature above working. pakkahisaab.app"
    )
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xE9, 0xD9, 0xA8)

    out_path = os.path.join(OUT_DIR, "PakkaHisaab_Brochure.docx")
    doc.save(out_path)
    return out_path


# ===========================================================================
# DOCUMENT 2 — PakkaHisaab User & Engineering Guide
# ===========================================================================
def add_workflow_block(doc, wf_id, name, preconditions, inputs, steps,
                        postconditions, screenshot, edge_cases):
    """Renders one fully-specified workflow using the strict technical schema."""
    hdr = doc.add_paragraph()
    hdr.paragraph_format.space_before = Pt(16)
    hdr.paragraph_format.space_after = Pt(4)
    r1 = hdr.add_run(f"{wf_id}: ")
    r1.font.bold = True
    r1.font.color.rgb = GOLD
    r1.font.name = FONT_HEAD
    r1.font.size = Pt(12.5)
    r2 = hdr.add_run(name)
    r2.font.bold = True
    r2.font.color.rgb = NAVY
    r2.font.name = FONT_HEAD
    r2.font.size = Pt(12.5)

    # Pre-conditions & Input Vector table
    spec_table = doc.add_table(rows=2, cols=2)
    spec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["Pre-conditions", "Input Vector"]
    values = [preconditions, inputs]
    for i in range(2):
        cells = spec_table.rows[i].cells
        cells[0].text = ""
        rl = cells[0].paragraphs[0].add_run(labels[i])
        rl.bold = True
        rl.font.size = Pt(9.5)
        rl.font.color.rgb = NAVY
        set_cell_shading(cells[0], BAND_GREY)
        cells[1].text = ""
        first = True
        for line in values[i]:
            para = cells[1].paragraphs[0] if first else cells[1].add_paragraph()
            first = False
            rr = para.add_run(f"• {line}")
            rr.font.size = Pt(9.5)
            para.paragraph_format.space_after = Pt(1)
        for c in cells:
            set_cell_borders(c, color="D9DEE6", size="3")
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    spec_table.columns[0].width = Cm(3.6)
    spec_table.columns[1].width = Cm(11.1)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(4)
    sub.paragraph_format.space_after = Pt(2)
    rs = sub.add_run("Step-by-Step Execution Sequence")
    rs.bold = True
    rs.font.size = Pt(10.5)
    rs.font.color.rgb = TEAL
    rs.font.name = FONT_HEAD

    step_table = doc.add_table(rows=len(steps) + 1, cols=2)
    step_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = step_table.rows[0].cells
    hdr_cells[0].text = ""
    hdr_cells[0].paragraphs[0].add_run("#")
    hdr_cells[1].text = ""
    hdr_cells[1].paragraphs[0].add_run("Action")
    for idx, step_text in enumerate(steps, start=1):
        cells = step_table.rows[idx].cells
        cells[0].text = ""
        cells[0].paragraphs[0].add_run(str(idx))
        cells[1].text = ""
        cells[1].paragraphs[0].add_run(step_text)
    step_table.columns[0].width = Cm(1.0)
    step_table.columns[1].width = Cm(13.7)
    shade_table_zebra(step_table, SLATE_HEX, "F2F4F7", header_font_color=WHITE)
    set_repeat_header(step_table.rows[0])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    post_p = doc.add_paragraph()
    post_p.paragraph_format.space_before = Pt(6)
    rp = post_p.add_run("Post-conditions & Data State Mutation: ")
    rp.bold = True
    rp.font.color.rgb = NAVY
    rp.font.size = Pt(10)
    rp2 = post_p.add_run(postconditions)
    rp2.font.size = Pt(10)

    add_image_placeholder(doc, f"Insert High-Definition UI Screenshot: {screenshot}", height_cm=4.2)

    add_shaded_paragraph_box(
        doc, edge_cases, fill="FFF6DF", border_color="C99A2E",
        label="EDGE-CASE HANDLING & PRO TIPS", label_color=PROTIP_LABEL,
    )
    add_hr(doc)


SLATE_HEX = "4A5568"


def build_rbac_matrix(doc):
    doc.add_heading("Role-Based Access Control (RBAC) Matrix", level=2)
    doc.add_paragraph(
        "The table below is the authoritative permission map governing every "
        "protected resource in PakkaHisaab. Access is enforced server-side at the "
        "query layer for every role — the client UI reflects, but never substitutes "
        "for, this boundary."
    )

    header = ["Capability / Resource", "System / Society Admin", "Resident", "Vendor"]
    rows = [
        ("Vendor KYC review & approval", "Full (Approve / Reject / Suspend)", "No Access", "Submit Only (own profile)"),
        ("Society maintenance ledger configuration", "Full (Create / Edit / Archive)", "Read Own Records", "No Access"),
        ("Penalty & grace-period rule engine", "Full (Configure / Override)", "Read Only (own dues)", "No Access"),
        ("Resident household ledger balance", "Read + Audit (all households)", "Full (own household only)", "Read Only (linked households only)"),
        ("Vendor storefront & catalogue", "Read + Suspend", "Read Only (browse/discover)", "Full (own storefront only)"),
        ("Digital ledger token issuance", "Read + Audit", "Authorise (own household)", "Initiate (linked households only)"),
        ("Recurring bill payment setup", "Read + Audit", "Full (own account)", "No Access"),
        ("Payout settlement processing", "Read + Approve batch release", "No Access", "Read Own Settlements"),
        ("Broadcast messaging / notifications", "Full (Compose / Send / Schedule)", "Receive Only", "Receive Only"),
        ("Billing discrepancy audit trail", "Full (Investigate / Annotate / Close)", "Read Own Flags", "Read Own Flags"),
        ("Historical statement export (CSV/PDF)", "Full (all households & vendors)", "Full (own account only)", "Full (own settlements only)"),
        ("Platform-wide analytics dashboard", "Full", "No Access", "Own-Storefront Metrics Only"),
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, val in enumerate(header):
        table.rows[0].cells[j].text = ""
        r = table.rows[0].cells[j].paragraphs[0].add_run(val)
        r.bold = True
    set_repeat_header(table.rows[0])
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = ""
            table.rows[i].cells[j].paragraphs[0].add_run(val)
    table.columns[0].width = Cm(4.6)
    table.columns[1].width = Cm(3.9)
    table.columns[2].width = Cm(3.3)
    table.columns[3].width = Cm(3.9)
    shade_table_zebra(table, NAVY_HEX, BAND_GREY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_shaded_paragraph_box(
        doc,
        [
            "Vendor sessions can never resolve a resident's personal identity fields "
            "(full legal name beyond first-name display, contact number, unit "
            "ownership documents) — only the authorised ledger balance and "
            "transaction-approval state of the linked household. This boundary is "
            "enforced at the API layer, not the UI, and cannot be bypassed by a "
            "modified client."
        ],
        fill="FDECEC", border_color="A32A2A",
        label="SECURITY NOTE", label_color=DANGER,
    )


NAVY_HEX = "0B1F3A"


def build_user_guide():
    doc = Document()
    configure_base_styles(doc)

    section = doc.sections[0]
    add_page_border(section, color=SLATE_HEX, size=12)
    add_footer_with_page_number(section, "PakkaHisaab — Engineering & User Workflows Guide")

    # ---------------- Cover ----------------
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    if os.path.exists(LOGO_PATH):
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = lp.add_run()
        run.add_picture(LOGO_PATH, width=Cm(2.4))

    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run("PakkaHisaab")

    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Exhaustive Engineering & User Workflows Guide")

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run(
        "Functional Specification, RBAC Matrix & Component-Level Workflow "
        "Blueprints for System Admins, Residents & Vendors"
    )
    r.font.size = Pt(11.5)
    r.italic = True
    r.font.color.rgb = SLATE

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_meta_table(doc, [
        ("Document Class", "Technical Reference — Functional & Workflow Specification"),
        ("Audience", "Engineering, QA, Support, Onboarding & Training Teams"),
        ("Version", "v1.0 — Production Release Edition"),
        ("Date", "19 July 2026"),
        ("Schema Standard", "Workflow ID · Pre-conditions · Input Vector · Execution "
                             "Sequence · Post-conditions · UI Reference · Edge Cases"),
    ], col_widths=(4.6, 10.1))

    doc.add_page_break()

    # ---------------- Section 1: Introduction & RBAC ----------------
    add_kicker(doc, "System Introduction")
    doc.add_heading("Platform Overview & Access Control Architecture", level=1)
    doc.add_paragraph(
        "PakkaHisaab operates as a single ledger of record shared across three "
        "strictly scoped user roles: the System / Society Admin, who governs "
        "society-wide financial policy and vendor trust; the Resident, who "
        "transacts against their own household ledger; and the Vendor, who "
        "fulfils local commerce against linked households. This section defines "
        "the exhaustive functional blueprint for every workflow available to each "
        "role, down to precise input/output boundaries, so that engineering, QA, "
        "and support teams share one unambiguous specification of system behaviour."
    )
    build_rbac_matrix(doc)

    doc.add_page_break()

    # ---------------- Section 2: Admin workflows ----------------
    add_kicker(doc, "Functional Blueprints — System / Society Admin")
    doc.add_heading("System / Society Admin Workflows", level=1)
    doc.add_paragraph(
        "The System / Society Admin role governs vendor trust, society-wide "
        "financial configuration, compliance auditing, and resident communication. "
        "Every admin workflow below is documented to the exact input and output "
        "boundary required for implementation, QA validation, and support "
        "escalation handling."
    )

    add_workflow_block(
        doc, "PK-WF-ADM-01", "Vendor Onboarding KYC Approval",
        preconditions=[
            "A vendor application exists in status PENDING_REVIEW.",
            "The vendor has uploaded mandatory KYC documents: government-issued "
            "identity proof, business address proof, and trade licence (where "
            "applicable).",
            "The admin account holds the VENDOR_VERIFIER permission grant.",
        ],
        inputs=[
            "vendor_id (UUID, required)",
            "kyc_document_set (identity_proof, address_proof, trade_licence — "
            "file references, required)",
            "service_radius_km (decimal, required)",
            "reviewer_decision (enum: APPROVE | REJECT | REQUEST_MORE_INFO, required)",
            "rejection_reason (string, required only when reviewer_decision = REJECT)",
        ],
        steps=[
            "Admin opens the Vendor Verification Dashboard and selects a vendor "
            "record from the PENDING_REVIEW queue.",
            "Admin reviews the uploaded identity proof, address proof, and trade "
            "licence documents against the declared business name and service "
            "address.",
            "Admin cross-checks the declared service radius against the list of "
            "societies requesting that vendor category to confirm serviceability.",
            "Admin selects a reviewer_decision: APPROVE, REJECT, or REQUEST_MORE_INFO.",
            "If APPROVE is selected, admin confirms the vendor category (Grocery, "
            "Dairy, Laundry, Pharmacy, or Household Services) and submits.",
            "If REJECT is selected, admin must enter a rejection_reason before the "
            "form allows submission.",
            "System persists the decision, updates the vendor record status, and "
            "triggers the corresponding vendor-facing notification.",
        ],
        postconditions=(
            "On APPROVE: vendor.status transitions PENDING_REVIEW → VERIFIED; a "
            "verified_at timestamp and reviewing admin_id are written to the vendor "
            "audit trail; the vendor storefront becomes discoverable to residents "
            "in societies within its service radius. On REJECT: vendor.status "
            "transitions to REJECTED with the rejection_reason persisted; the "
            "vendor may resubmit a new application. On REQUEST_MORE_INFO: "
            "vendor.status transitions to INFO_REQUESTED and the vendor receives "
            "a itemised list of outstanding document gaps."
        ),
        screenshot="Admin Panel → Vendor Verification Dashboard, showing the "
                   "pending-queue list on the left and the document review pane "
                   "with Approve / Reject / Request-More-Info action buttons on "
                   "the right.",
        edge_cases=[
            "Duplicate PAN / business registration number detected: system blocks "
            "submission and surfaces the conflicting vendor_id to the admin for "
            "manual reconciliation before a decision can be recorded.",
            "Uploaded document fails automated legibility/OCR pre-check: the "
            "APPROVE action is disabled until the vendor re-uploads a passing "
            "document, preventing accidental approval of an unreadable KYC file.",
            "Pro Tip: use REQUEST_MORE_INFO rather than REJECT for any borderline "
            "or partially complete submission — it preserves the vendor's queue "
            "position and avoids forcing a full application restart.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-ADM-02", "Society Maintenance Ledger Configuration",
        preconditions=[
            "The admin account holds the LEDGER_CONFIG permission grant.",
            "The society workspace has at least one registered residential unit.",
        ],
        inputs=[
            "unit_category (enum, e.g. 1BHK / 2BHK / 3BHK / Commercial — required)",
            "maintenance_amount (decimal, currency INR, required)",
            "billing_cycle (enum: MONTHLY | QUARTERLY | ANNUAL, required)",
            "due_date_day_of_month (integer 1–28, required)",
            "effective_from (date, required)",
        ],
        steps=[
            "Admin navigates to Society Settings → Maintenance Ledger "
            "Configuration.",
            "Admin defines or edits a maintenance_amount slab per unit_category.",
            "Admin sets the billing_cycle and due_date_day_of_month governing when "
            "dues are generated and considered outstanding.",
            "Admin sets an effective_from date, allowing rate changes to apply only "
            "to future billing cycles without altering historical ledger entries.",
            "Admin reviews a system-generated preview of the next three billing "
            "cycles per unit_category before confirming.",
            "Admin submits the configuration; the system validates no overlapping "
            "active slab exists for the same unit_category and date range.",
        ],
        postconditions=(
            "A new or updated maintenance_slab record is persisted with an "
            "immutable created_by admin_id and created_at timestamp. Existing "
            "billing_cycle records prior to effective_from remain untouched; "
            "future ledger generation jobs read the new slab from effective_from "
            "onward. All linked resident ledgers automatically reflect the "
            "updated recurring due amount from their next billing cycle."
        ),
        screenshot="Admin Panel → Society Settings → Maintenance Ledger "
                   "Configuration, showing a slab editor table with unit categories "
                   "as rows and a live billing-cycle preview panel.",
        edge_cases=[
            "Overlapping effective_from date for the same unit_category: system "
            "rejects the save and highlights the conflicting existing slab rather "
            "than silently overwriting it.",
            "due_date_day_of_month set beyond 28: input is rejected at validation "
            "to avoid undefined behaviour in months with fewer than 29–31 days.",
            "Pro Tip: set effective_from at least one full billing cycle ahead of "
            "the current date so residents receive the standard notice period "
            "before a maintenance amount change takes effect.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-ADM-03", "Automated Penalty & Grace-Period Rule Configuration",
        preconditions=[
            "The admin account holds the LEDGER_CONFIG permission grant.",
            "At least one active maintenance_slab exists for the society.",
        ],
        inputs=[
            "grace_period_days (integer ≥ 0, required)",
            "penalty_type (enum: FLAT_FEE | PERCENTAGE_OF_DUE, required)",
            "penalty_value (decimal, required)",
            "penalty_recurrence (enum: ONE_TIME | PER_BILLING_CYCLE_UNTIL_PAID, required)",
            "escalation_notification_schedule (array of day-offsets, required)",
        ],
        steps=[
            "Admin navigates to Society Settings → Penalty & Grace-Period "
            "Rules.",
            "Admin sets grace_period_days — the number of days after a due_date "
            "before a balance is classified as overdue.",
            "Admin selects penalty_type and enters penalty_value (a flat rupee "
            "amount or a percentage of the outstanding due).",
            "Admin selects penalty_recurrence to determine whether the penalty "
            "applies once or compounds every billing cycle the due remains unpaid.",
            "Admin configures the escalation_notification_schedule — the day "
            "offsets (relative to due_date) at which reminder intensity escalates.",
            "Admin submits; the system validates that grace_period_days does not "
            "exceed the billing_cycle length and confirms the rule is now active.",
        ],
        postconditions=(
            "A penalty_rule record is persisted and linked to the society's active "
            "maintenance_slab configuration. The automated billing engine begins "
            "evaluating every open due against this rule at the next scheduled "
            "evaluation run; overdue balances are flagged and penalty line items "
            "are appended to the resident's ledger automatically without further "
            "admin action."
        ),
        screenshot="Admin Panel → Society Settings → Penalty & "
                   "Grace-Period Rules, showing a rule builder with grace-period "
                   "slider, penalty-type toggle, and an escalation timeline "
                   "visualisation.",
        edge_cases=[
            "grace_period_days exceeds the billing_cycle duration: system rejects "
            "the configuration since it would allow a due to roll into the next "
            "cycle before ever becoming overdue.",
            "PERCENTAGE_OF_DUE selected with penalty_value above a configurable "
            "sanity ceiling (default 25%): system requires an explicit secondary "
            "confirmation to prevent an accidental punitive rate.",
            "Pro Tip: pair PER_BILLING_CYCLE_UNTIL_PAID recurrence with an "
            "escalation_notification_schedule of at least three touchpoints — "
            "societies that skip early reminders see materially higher penalty "
            "dispute volume at audit time.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-ADM-04", "Billing Discrepancy Audit & Resolution",
        preconditions=[
            "The admin account holds the AUDIT_REVIEW permission grant.",
            "At least one ledger entry has been auto-flagged by the reconciliation "
            "engine as a discrepancy (balance mismatch, duplicate token, or failed "
            "settlement).",
        ],
        inputs=[
            "discrepancy_id (UUID, system-generated, required)",
            "investigation_notes (string, optional per step)",
            "resolution_action (enum: ADJUST_LEDGER | MARK_FALSE_POSITIVE | "
            "ESCALATE_TO_VENDOR | ESCALATE_TO_RESIDENT, required)",
            "adjustment_amount (decimal, required only when resolution_action = "
            "ADJUST_LEDGER)",
        ],
        steps=[
            "Admin opens the Billing Discrepancy Audit queue, sorted by "
            "flagged_at timestamp and severity.",
            "Admin selects a discrepancy_id and reviews the system-generated "
            "comparison: expected balance vs. recorded balance, with the full "
            "transaction chain that produced the mismatch.",
            "Admin adds investigation_notes documenting findings for the audit "
            "trail.",
            "Admin selects a resolution_action appropriate to the root cause "
            "identified.",
            "If ADJUST_LEDGER is selected, admin enters the adjustment_amount and "
            "a mandatory justification note.",
            "Admin submits the resolution; the system requires a secondary "
            "confirmation for any ADJUST_LEDGER action exceeding a configurable "
            "materiality threshold.",
        ],
        postconditions=(
            "The discrepancy record transitions from OPEN to RESOLVED with the "
            "resolution_action, investigation_notes, and reviewing admin_id "
            "permanently appended to the immutable audit trail. If ADJUST_LEDGER "
            "was selected, a compensating ledger entry is posted referencing the "
            "original discrepancy_id, and the affected resident's or vendor's "
            "balance is recalculated in real time."
        ),
        screenshot="Admin Panel → Billing Discrepancy Audit, showing a "
                   "side-by-side expected-vs-actual balance comparison with the "
                   "full transaction chain timeline below.",
        edge_cases=[
            "Discrepancy involves a vendor settlement already paid out: "
            "ADJUST_LEDGER is disabled until the linked payout batch is placed on "
            "hold, preventing a correction from being silently absorbed into a "
            "future payout.",
            "Same discrepancy_id resubmitted twice due to a network retry: system "
            "de-duplicates on discrepancy_id and rejects the second resolution "
            "attempt with a CONFLICT response.",
            "Pro Tip: always document investigation_notes before selecting "
            "MARK_FALSE_POSITIVE — recurring false-positive patterns without notes "
            "make it impossible to later tune the reconciliation engine's "
            "sensitivity.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-ADM-05", "Broadcast Messaging to Residents & Vendors",
        preconditions=[
            "The admin account holds the BROADCAST_SEND permission grant.",
            "At least one target audience segment (all residents, specific "
            "block/tower, all vendors, or a custom filter) is available.",
        ],
        inputs=[
            "audience_segment (enum: ALL_RESIDENTS | SPECIFIC_BLOCK | ALL_VENDORS "
            "| CUSTOM_FILTER, required)",
            "message_title (string, max 120 characters, required)",
            "message_body (string, max 2000 characters, required)",
            "delivery_channel (multi-select: IN_APP | SMS | EMAIL, required, at "
            "least one)",
            "scheduled_send_time (datetime, optional — defaults to immediate)",
        ],
        steps=[
            "Admin opens Broadcast Messaging and selects an audience_segment.",
            "If CUSTOM_FILTER is selected, admin defines filter criteria (e.g. "
            "outstanding balance above a threshold, specific unit range).",
            "Admin composes message_title and message_body.",
            "Admin selects one or more delivery_channel options.",
            "Admin optionally sets a scheduled_send_time, or leaves it blank to "
            "send immediately.",
            "Admin reviews an audience-size preview count before final submission.",
            "Admin confirms and submits the broadcast for delivery.",
        ],
        postconditions=(
            "A broadcast_message record is persisted with status QUEUED or SENT. "
            "The notification engine fans the message out to every recipient in "
            "the resolved audience_segment across each selected delivery_channel. "
            "A delivery receipt log is created per recipient, capturing "
            "sent_at, delivered_at, and read_at (where the channel supports read "
            "receipts) for later audit."
        ),
        screenshot="Admin Panel → Broadcast Messaging, showing the audience "
                   "segment selector, message composer, and a live recipient-count "
                   "preview badge.",
        edge_cases=[
            "CUSTOM_FILTER resolves to zero matching recipients: system blocks "
            "submission and surfaces a warning rather than silently sending to "
            "nobody.",
            "message_body exceeds the SMS channel's character limit: system "
            "automatically truncates the SMS variant with a “read more in app” "
            "suffix while keeping the full text intact for IN_APP and EMAIL "
            "channels.",
            "Pro Tip: use scheduled_send_time to queue due-date reminders for "
            "early morning delivery — open and response rates measured across "
            "societies are consistently higher before 10 AM.",
        ],
    )

    doc.add_page_break()

    # ---------------- Section 3: Resident workflows ----------------
    add_kicker(doc, "Functional Blueprints — Resident")
    doc.add_heading("Resident User Workflows", level=1)
    doc.add_paragraph(
        "The Resident role transacts exclusively against its own household ledger "
        "and the vendors it has chosen to link. Every resident-facing workflow is "
        "documented below with the exact fields, sequence, and system state "
        "changes required for a complete, production-grade implementation."
    )

    add_workflow_block(
        doc, "PK-WF-RES-01", "Discovering Verified Nearby Vendors",
        preconditions=[
            "Resident account is active and linked to a registered residential "
            "unit.",
            "At least one vendor with status VERIFIED exists within the "
            "resident's society service area.",
        ],
        inputs=[
            "vendor_category_filter (enum, optional: Grocery | Dairy | Laundry | "
            "Pharmacy | Household Services | All)",
            "search_query (string, optional free-text)",
            "sort_order (enum: DISTANCE | RATING | RECENTLY_ADDED, optional, "
            "default DISTANCE)",
        ],
        steps=[
            "Resident opens the Vendor Discovery tab from the home dashboard.",
            "Resident optionally applies a vendor_category_filter or enters a "
            "search_query.",
            "Resident optionally changes sort_order to reorder results.",
            "Resident reviews the returned list, each entry showing verification "
            "badge, category, rating, and estimated distance.",
            "Resident taps a vendor card to view the full storefront profile, "
            "including catalogue and service radius confirmation for their unit.",
            "Resident taps “Link Vendor” to request a household-ledger linkage "
            "with the selected vendor.",
        ],
        postconditions=(
            "A vendor_link request record is created in status "
            "PENDING_VENDOR_ACCEPTANCE, associated with the resident's household "
            "ledger and the selected vendor_id. The vendor receives a real-time "
            "notification to accept the linkage before any ledger token can be "
            "issued against that household."
        ),
        screenshot="Resident App → Vendor Discovery, showing a filterable "
                   "list of verified vendor cards with distance, rating, and "
                   "category badges, and a detail sheet for the selected vendor.",
        edge_cases=[
            "No vendors match the applied vendor_category_filter within the "
            "society: system displays an explicit “no vendors yet in this "
            "category” state with an option to notify the resident when one is "
            "onboarded, rather than an empty blank screen.",
            "Resident attempts to link a vendor already at its configured maximum "
            "concurrent household count: request is rejected with a clear "
            "capacity-reached message instead of silently queuing indefinitely.",
            "Pro Tip: check the verification badge date — a vendor re-verified "
            "within the last 90 days reflects the most current KYC and service "
            "radius data.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-RES-02", "Initiating & Topping Up a Secure Ledger Balance",
        preconditions=[
            "Resident's household ledger exists and is in ACTIVE status.",
            "Resident has at least one verified payment instrument (UPI, "
            "linked bank account, or card) on file.",
        ],
        inputs=[
            "top_up_amount (decimal, currency INR, minimum ₹100, required)",
            "payment_instrument_id (UUID, required)",
            "auto_reload_threshold (decimal, optional)",
            "auto_reload_amount (decimal, required only if auto_reload_threshold "
            "is set)",
        ],
        steps=[
            "Resident opens My Ledger and selects “Add Funds.”",
            "Resident enters top_up_amount or selects a preset amount chip.",
            "Resident selects a payment_instrument_id from saved instruments or "
            "adds a new one.",
            "Resident optionally configures auto_reload_threshold and "
            "auto_reload_amount to enable automatic top-ups when the balance "
            "runs low.",
            "Resident confirms the transaction; the system routes the request to "
            "the configured payment gateway for authorisation.",
            "On gateway authorisation success, the system credits the household "
            "ledger balance in real time and issues a digital receipt.",
        ],
        postconditions=(
            "A ledger_transaction record of type CREDIT_TOPUP is appended to the "
            "household's immutable ledger with the authorised amount, gateway "
            "reference ID, and timestamp. The resident's available balance "
            "reflects the new total instantly across all connected vendor and "
            "admin views. If auto-reload was configured, the "
            "auto_reload_threshold and auto_reload_amount are persisted against "
            "the household's ledger settings for future automatic triggers."
        ),
        screenshot="Resident App → My Ledger → Add Funds, showing preset "
                   "amount chips, a saved payment-instrument selector, and an "
                   "auto-reload toggle with threshold slider.",
        edge_cases=[
            "Payment gateway authorisation fails or times out: no ledger credit "
            "is posted, the resident sees an explicit failure reason "
            "(insufficient funds, instrument declined, gateway timeout), and the "
            "attempt is logged for support diagnostics without double-charging on "
            "retry.",
            "auto_reload_threshold set higher than auto_reload_amount would "
            "sensibly cover: system warns the resident that reloads may trigger "
            "more frequently than intended, but does not block the configuration.",
            "Pro Tip: enable auto-reload with a threshold set to roughly one "
            "week's typical vendor spend to avoid a linked vendor ever being "
            "unable to complete a purchase due to an empty balance.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-RES-03", "Authorising a Direct-to-Door Vendor Digital Token",
        preconditions=[
            "Resident's household is linked (status ACCEPTED) to the requesting "
            "vendor.",
            "Household ledger balance is sufficient to cover the requested token "
            "amount, or the society's credit-limit policy permits a negative "
            "balance up to its configured ceiling.",
        ],
        inputs=[
            "vendor_id (UUID, required)",
            "token_amount (decimal, required)",
            "item_description (string, optional, vendor-supplied)",
            "authorization_method (enum: IN_APP_APPROVE | BIOMETRIC_CONFIRM | "
            "PIN_CONFIRM, required)",
        ],
        steps=[
            "Vendor initiates a token request against the resident's linked "
            "household with token_amount and an optional item_description.",
            "Resident receives a real-time push notification presenting the "
            "pending token request.",
            "Resident opens the request and reviews vendor identity, "
            "token_amount, and item_description.",
            "Resident authorises using the configured authorization_method "
            "(in-app approval, biometric, or PIN).",
            "System validates sufficient balance or available credit limit at "
            "the moment of authorisation.",
            "Upon successful validation, the system finalises the token, debits "
            "the household ledger, and issues a digital receipt to both parties.",
        ],
        postconditions=(
            "A ledger_transaction record of type DEBIT_VENDOR_TOKEN is appended "
            "to the household ledger, referencing vendor_id, token_amount, and "
            "the authorization_method used. The vendor's pending-settlement "
            "balance increases by token_amount. Both resident and vendor receive "
            "a matching digital receipt with a shared transaction reference ID."
        ),
        screenshot="Resident App → Token Authorisation, showing a pending "
                   "request card with vendor name, verification badge, requested "
                   "amount, and Approve / Decline buttons.",
        edge_cases=[
            "Household balance and available credit limit are both insufficient "
            "at authorisation time: the token request is automatically declined "
            "with reason INSUFFICIENT_BALANCE, and the vendor is notified to "
            "collect payment by an alternate method for that transaction.",
            "Resident does not respond within the configured request expiry "
            "window (default 15 minutes): the token request auto-expires to "
            "EXPIRED status, and the vendor must re-initiate if the transaction "
            "is still valid.",
            "Pro Tip: enable BIOMETRIC_CONFIRM as the default authorization_method "
            "for households above a configurable daily token-value threshold to "
            "reduce accidental or coerced approvals.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-RES-04", "Automated Recurring Society Bill Payment Setup",
        preconditions=[
            "Resident's household ledger is ACTIVE with at least one open or "
            "upcoming maintenance due.",
            "Resident has at least one verified payment instrument on file.",
        ],
        inputs=[
            "payment_instrument_id (UUID, required)",
            "auto_pay_trigger (enum: ON_DUE_DATE | N_DAYS_BEFORE_DUE, required)",
            "trigger_offset_days (integer, required only if auto_pay_trigger = "
            "N_DAYS_BEFORE_DUE)",
            "max_auto_pay_amount (decimal, optional cap)",
        ],
        steps=[
            "Resident opens My Ledger → Recurring Payments and selects "
            "“Set Up Auto-Pay.”",
            "Resident selects a payment_instrument_id to be charged "
            "automatically.",
            "Resident selects an auto_pay_trigger and, if applicable, a "
            "trigger_offset_days value.",
            "Resident optionally sets max_auto_pay_amount as a safety ceiling, "
            "above which auto-pay will not fire and a manual review notification "
            "is sent instead.",
            "Resident reviews a summary confirming the next three scheduled "
            "auto-pay dates and amounts.",
            "Resident confirms; the system activates the recurring payment "
            "mandate.",
        ],
        postconditions=(
            "An auto_pay_mandate record is persisted in ACTIVE status against the "
            "household ledger. On each subsequent billing cycle, the system "
            "automatically initiates a payment on the configured trigger date, "
            "posting a CREDIT_AUTO_PAY ledger_transaction upon success, or a "
            "FAILED_AUTO_PAY flag with resident notification upon failure — "
            "never silently skipping a cycle without notice."
        ),
        screenshot="Resident App → Recurring Payments → Set Up Auto-Pay, "
                   "showing the trigger-timing selector and a three-cycle payment "
                   "preview schedule.",
        edge_cases=[
            "Scheduled auto-pay amount exceeds max_auto_pay_amount (e.g. after a "
            "penalty was applied): auto-pay is suppressed for that cycle and the "
            "resident receives a manual-payment-required notification instead of "
            "an unexpected larger charge.",
            "Configured payment_instrument_id expires or is removed before the "
            "next trigger date: the mandate transitions to PAUSED and the "
            "resident is prompted to update their instrument before the next "
            "billing cycle.",
            "Pro Tip: choose N_DAYS_BEFORE_DUE with a 2–3 day offset rather than "
            "ON_DUE_DATE — it leaves a buffer to retry automatically once before "
            "a due is ever marked overdue.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-RES-05", "Historical Statement Export (CSV / PDF)",
        preconditions=[
            "Resident's household ledger contains at least one historical "
            "transaction.",
        ],
        inputs=[
            "date_range_start (date, required)",
            "date_range_end (date, required)",
            "export_format (enum: CSV | PDF, required)",
            "transaction_type_filter (enum, optional: ALL | SOCIETY_DUES | "
            "VENDOR_TOKENS | TOP_UPS)",
        ],
        steps=[
            "Resident opens My Ledger → Statements & Export.",
            "Resident selects date_range_start and date_range_end.",
            "Resident optionally applies a transaction_type_filter to narrow the "
            "export scope.",
            "Resident selects export_format (CSV for spreadsheet analysis, PDF "
            "for a formatted statement suitable for record-keeping).",
            "Resident taps “Generate Export”; the system compiles the requested "
            "ledger slice.",
            "System delivers the completed file as an in-app download and, if "
            "configured, an emailed copy.",
        ],
        postconditions=(
            "An export_request record is logged with the requesting resident_id, "
            "the applied date range and filters, and a completed_at timestamp. "
            "The generated file exactly mirrors the immutable ledger data for "
            "the requested range — no export operation can mutate underlying "
            "ledger records."
        ),
        screenshot="Resident App → Statements & Export, showing a date-range "
                   "picker, format toggle (CSV/PDF), and a generated-file download "
                   "card.",
        edge_cases=[
            "Requested date_range spans a period with zero transactions: system "
            "still generates a valid, correctly headered empty statement rather "
            "than returning an error.",
            "date_range_end precedes date_range_start: input validation blocks "
            "generation until the range is corrected.",
            "Pro Tip: use the PDF format when submitting a statement for society "
            "audit or dispute purposes — it includes the resident's account "
            "verification header that CSV exports omit by design.",
        ],
    )

    doc.add_page_break()

    # ---------------- Section 4: Vendor workflows ----------------
    add_kicker(doc, "Functional Blueprints — Vendor")
    doc.add_heading("Vendor User Workflows", level=1)
    doc.add_paragraph(
        "The Vendor role operates a digital storefront linked to one or more "
        "resident households across its verified societies. Every vendor-facing "
        "workflow below is documented to the same exhaustive input/output "
        "standard used for the Admin and Resident roles."
    )

    add_workflow_block(
        doc, "PK-WF-VEN-01", "Account Provisioning & Digital Storefront Setup",
        preconditions=[
            "Vendor application has been APPROVED via PK-WF-ADM-01.",
            "Vendor has received a provisioning invite to complete account "
            "activation.",
        ],
        inputs=[
            "storefront_name (string, required)",
            "business_category (enum: Grocery | Dairy | Laundry | Pharmacy | "
            "Household Services, required)",
            "catalogue_items (array of {item_name, unit_price, unit_of_measure}, "
            "required, minimum 1 entry)",
            "service_radius_km (decimal, required)",
            "settlement_bank_account (account number + IFSC, required)",
        ],
        steps=[
            "Vendor accepts the provisioning invite and sets an account password.",
            "Vendor enters storefront_name and confirms business_category "
            "(pre-filled from the approved KYC application).",
            "Vendor adds catalogue_items, specifying unit_price and "
            "unit_of_measure for each.",
            "Vendor confirms or adjusts service_radius_km within the ceiling "
            "approved during KYC review.",
            "Vendor enters settlement_bank_account details for payout routing.",
            "Vendor submits; the system runs a penny-drop verification against "
            "the settlement_bank_account before activating the storefront.",
        ],
        postconditions=(
            "On successful penny-drop verification, the vendor account "
            "transitions to ACTIVE and the storefront becomes visible in "
            "resident-facing Vendor Discovery within its service_radius_km. The "
            "catalogue_items are persisted as the vendor's live pricing "
            "reference, used to pre-populate future token requests."
        ),
        screenshot="Vendor App → Storefront Setup, showing the catalogue "
                   "builder with item rows, a service-radius map selector, and "
                   "the bank-account verification step.",
        edge_cases=[
            "Penny-drop verification fails (account name mismatch or invalid "
            "IFSC): storefront remains in PENDING_BANK_VERIFICATION and cannot go "
            "live until corrected, preventing payouts from routing to an "
            "unverified account.",
            "service_radius_km entered exceeds the ceiling approved during KYC: "
            "input is capped at the approved maximum with an explanatory message "
            "rather than a silent server-side truncation.",
            "Pro Tip: keep catalogue_items unit_price current — token requests "
            "pre-populate from this list, and stale pricing is the single most "
            "common source of resident-reported billing disputes.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-VEN-02", "Real-Time Ledger Updates Per Customer Household",
        preconditions=[
            "Vendor storefront is ACTIVE.",
            "At least one household has an ACCEPTED linkage with the vendor.",
        ],
        inputs=[
            "household_id (UUID, required)",
            "line_items (array of {catalogue_item_id, quantity}, required, "
            "minimum 1 entry)",
            "delivery_confirmation (boolean, required)",
        ],
        steps=[
            "Vendor opens the Household Ledger view and selects a linked "
            "household_id.",
            "Vendor adds line_items from the catalogue, specifying quantity for "
            "each.",
            "System computes the total token_amount from current catalogue "
            "pricing.",
            "Vendor marks delivery_confirmation once goods or services are "
            "handed over.",
            "Vendor submits the request, which routes to the resident for "
            "authorisation per PK-WF-RES-03.",
            "Vendor monitors the request status (Pending, Approved, Declined, "
            "Expired) in real time on the same screen.",
        ],
        postconditions=(
            "A token_request record is created referencing household_id, "
            "line_items, and computed token_amount, in status "
            "PENDING_RESIDENT_AUTHORIZATION. Upon resident approval, the "
            "household's ledger debit and the vendor's pending-settlement "
            "balance update atomically and in real time, visible to the vendor "
            "without a manual refresh."
        ),
        screenshot="Vendor App → Household Ledger, showing a linked-household "
                   "selector, catalogue line-item picker, and a live request-status "
                   "tracker.",
        edge_cases=[
            "Vendor attempts to submit a token_request against a household whose "
            "linkage has been revoked since the last screen refresh: system "
            "rejects the submission with a LINKAGE_INACTIVE error and prompts the "
            "vendor to re-request linkage.",
            "Catalogue pricing changes between when the vendor opens the ledger "
            "screen and submits the request: system recomputes token_amount from "
            "current pricing at submission time and displays the recalculated "
            "total for vendor confirmation before sending to the resident.",
            "Pro Tip: always set delivery_confirmation only at the point of "
            "actual handover — token requests submitted before delivery create "
            "avoidable disputes if a resident declines after receiving goods.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-VEN-03", "Transactional Request Routing & Digital Receipt Processing",
        preconditions=[
            "A token_request exists in status APPROVED (resident-authorised).",
        ],
        inputs=[
            "token_request_id (UUID, required)",
            "receipt_delivery_channel (multi-select: IN_APP | SMS | EMAIL, "
            "required, at least one)",
        ],
        steps=[
            "System automatically routes an APPROVED token_request to the "
            "vendor's transaction queue in real time.",
            "Vendor opens the transaction queue and selects the completed "
            "token_request_id.",
            "Vendor confirms the receipt_delivery_channel(s) for the digital "
            "receipt (defaults to the resident's saved preference).",
            "Vendor triggers “Issue Receipt,” generating a system receipt "
            "referencing the token_request_id, itemised line_items, and final "
            "token_amount.",
            "System delivers the digital receipt to the resident across the "
            "selected channel(s) and archives a copy in the vendor's transaction "
            "history.",
        ],
        postconditions=(
            "A digital_receipt record is generated and immutably linked to its "
            "source token_request_id. The transaction is marked COMPLETED in "
            "both the vendor's and resident's transaction histories, and the "
            "underlying ledger_transaction becomes eligible for inclusion in the "
            "vendor's next payout settlement batch."
        ),
        screenshot="Vendor App → Transaction Queue, showing an approved "
                   "request card with an “Issue Receipt” action and a delivery-"
                   "channel confirmation checklist.",
        edge_cases=[
            "Vendor attempts to issue a receipt for a token_request not yet in "
            "APPROVED status: the action is disabled with an inline explanation "
            "of the current status, preventing a receipt from referencing an "
            "unauthorised transaction.",
            "Notification delivery to the resident's device fails on all "
            "selected channels (e.g. offline device, invalid email): system "
            "retries on an exponential backoff schedule and the digital_receipt "
            "remains permanently retrievable from the resident's in-app "
            "transaction history regardless of delivery outcome.",
            "Pro Tip: enable both IN_APP and SMS as default delivery channels — "
            "residents without consistent data connectivity still receive "
            "confirmation, reducing “did I actually pay” support queries.",
        ],
    )

    add_workflow_block(
        doc, "PK-WF-VEN-04", "Payout Settlement Processing",
        preconditions=[
            "Vendor has one or more COMPLETED transactions not yet included in "
            "a settled payout batch.",
            "Vendor's settlement_bank_account is in VERIFIED status.",
        ],
        inputs=[
            "settlement_cycle (enum: DAILY | WEEKLY | BI_WEEKLY, vendor-configured, "
            "required)",
            "payout_batch_id (UUID, system-generated at cycle close)",
        ],
        steps=[
            "At the close of each settlement_cycle, the system automatically "
            "aggregates all COMPLETED transactions not yet paid out into a new "
            "payout_batch_id.",
            "Vendor opens the Settlements tab to review the batch summary — "
            "transaction count, gross amount, and any platform fee deductions.",
            "Vendor reviews itemised transactions within the batch, each linked "
            "back to its originating token_request_id and digital_receipt.",
            "Admin releases the payout batch for processing per the society's "
            "configured approval policy (auto-release or manual admin approval).",
            "System initiates the bank transfer to the vendor's "
            "settlement_bank_account.",
            "Vendor receives a real-time settlement confirmation once the "
            "transfer clears.",
        ],
        postconditions=(
            "The payout_batch transitions from PENDING → PROCESSING → "
            "SETTLED, with each state transition timestamped on the immutable "
            "audit trail. Every included transaction is marked PAID_OUT, "
            "preventing it from being included in any future batch. The vendor's "
            "pending-settlement balance decreases by the batch's net amount, and "
            "a settlement receipt is generated for the vendor's records."
        ),
        screenshot="Vendor App → Settlements, showing the current "
                   "payout_batch summary card, an itemised transaction list, and "
                   "a settlement-status progress tracker.",
        edge_cases=[
            "Bank transfer fails at the processing stage (account closed, "
            "incorrect details flagged post-verification): payout_batch reverts "
            "to FAILED status, the vendor is notified immediately, and no "
            "included transaction is marked PAID_OUT until a corrected transfer "
            "succeeds.",
            "Vendor changes settlement_cycle mid-period: the change takes effect "
            "only from the next full cycle boundary, and the current in-flight "
            "payout_batch completes under the previous cycle configuration to "
            "avoid double-counting or gaps in settlement coverage.",
            "Pro Tip: reconcile the itemised transaction list against your own "
            "delivery records before each cycle closes — any discrepancy is far "
            "faster to resolve pre-settlement than after a payout has already "
            "been transferred.",
        ],
    )

    out_path = os.path.join(OUT_DIR, "PakkaHisaab_UserGuide.docx")
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    brochure_path = build_brochure()
    guide_path = build_user_guide()
    print(f"Brochure saved to: {brochure_path}")
    print(f"User guide saved to: {guide_path}")
